#!/usr/bin/env python3
"""
convert_grammar_guide.py
========================
One-shot conversion of ``data lake/Aelaki Grammar Guide.docx`` → markdown
so the megadoc's content is greppable alongside the other ``data lake/*.md``
reference material. The docx stays in place as the authoritative original.

Usage:
    python wiki-scripts/convert_grammar_guide.py
"""
import io
import sys
from pathlib import Path

# Windows console defaults to cp1252; force utf-8 on stdout for messages.
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.oxml.ns import qn

REPO = Path(__file__).resolve().parent.parent
SRC = REPO / "data lake" / "Aelaki Grammar Guide.docx"
DST = REPO / "data lake" / "Aelaki_Grammar_Guide.md"

HEADING_PREFIX = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
    "Heading 5": "##### ",
    "Heading 6": "###### ",
    "Title": "# ",
    "Subtitle": "## ",
}


def cell_text(cell) -> str:
    """Flatten a cell's paragraphs into a single whitespace-joined string."""
    parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    # Escape pipes so markdown table structure isn't broken.
    return " ".join(parts).replace("|", "\\|").replace("\n", " ")


def table_to_md(table) -> str:
    rows = table.rows
    if not rows:
        return ""
    header_cells = [cell_text(c) for c in rows[0].cells]
    width = len(header_cells)
    out = ["| " + " | ".join(header_cells) + " |",
           "|" + "|".join(["---"] * width) + "|"]
    for row in rows[1:]:
        cells = [cell_text(c) for c in row.cells]
        # Pad/trim to header width so markdown renders cleanly.
        cells = (cells + [""] * width)[:width]
        out.append("| " + " | ".join(cells) + " |")
    return "\n".join(out)


def render_paragraph(p) -> str:
    style = p.style.name if p.style else ""
    text = p.text.rstrip()
    if not text:
        return ""
    prefix = HEADING_PREFIX.get(style, "")
    if prefix:
        return prefix + text
    if style.startswith("List"):
        return f"- {text}"
    return text


def iter_body(doc):
    """Yield ('para', Paragraph) / ('table', Table) in document order."""
    from docx.oxml.ns import qn  # noqa: F401 already imported
    body = doc.element.body
    paras_by_elem = {p._element: p for p in doc.paragraphs}
    tables_by_elem = {t._element: t for t in doc.tables}
    for child in body.iterchildren():
        if child in paras_by_elem:
            yield "para", paras_by_elem[child]
        elif child in tables_by_elem:
            yield "table", tables_by_elem[child]


def main() -> int:
    if not SRC.is_file():
        print(f"ERROR: {SRC} not found", file=sys.stderr)
        return 1

    doc = Document(str(SRC))
    out_lines = [
        "<!-- Converted from 'Aelaki Grammar Guide.docx' by wiki-scripts/convert_grammar_guide.py.",
        "     The docx remains authoritative; this markdown is a greppable mirror. -->",
        "",
    ]

    blank = False
    for kind, node in iter_body(doc):
        if kind == "para":
            line = render_paragraph(node)
            if not line:
                if not blank:
                    out_lines.append("")
                    blank = True
                continue
            out_lines.append(line)
            out_lines.append("")
            blank = True
        else:  # table
            out_lines.append(table_to_md(node))
            out_lines.append("")
            blank = True

    DST.write_text("\n".join(out_lines), encoding="utf-8")
    print(f"Wrote {DST.relative_to(REPO)} "
          f"({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, "
          f"{DST.stat().st_size} bytes).")
    return 0


if __name__ == "__main__":
    sys.exit(main())
